#!/usr/bin/env python3
"""
Generate professional Esper-branded pandoc reference templates.

Each template is a .docx file that pandoc uses as --reference-doc to style output.
Pandoc reads ALL defined styles from the reference doc and applies them to converted content.

Brand: Esper (light touch)
- Navy: #1A237E (headings, accents)
- Blue: #0288D1 (links, subtle accents)  
- Text: #212121 (body — NOT pure black)
- Gray: #757575 (secondary text, footers)
- Light gray: #F5F5F5 (code backgrounds)
- Table header: #1A237E with white text
- Table alt rows: #F5F8FC

Fonts: Calibri (body), Calibri Light (headings), Consolas (code)

Usage:
    python3 generate_templates.py

Output directory resolution (in priority order):
    1. PANDOC_TEMPLATE_DIR environment variable, if set
    2. The skill's reference/ directory (when run from within the repo)
    3. Fallback: ~/.config/opencode/templates

Requires: python-docx  (pip install python-docx)
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# === Brand Colors ===
NAVY = RGBColor(0x1A, 0x23, 0x7E)
BLUE = RGBColor(0x02, 0x88, 0xD1)
TEXT_DARK = RGBColor(0x21, 0x21, 0x21)
TEXT_GRAY = RGBColor(0x75, 0x75, 0x75)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF5, 0xF8, 0xFC)
CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)
BORDER_GRAY = RGBColor(0xDD, 0xDD, 0xDD)
RULE_NAVY = "1A237E"


def resolve_output_dir() -> Path:
    """Resolve where to write templates.

    Priority:
      1. PANDOC_TEMPLATE_DIR env var
      2. <repo>/reference  (this script lives in <repo>/scripts/)
      3. ~/.config/opencode/templates
    """
    env_dir = os.environ.get("PANDOC_TEMPLATE_DIR")
    if env_dir:
        return Path(os.path.expanduser(env_dir))

    # This script is at <repo>/scripts/generate_templates.py
    repo_reference = Path(__file__).resolve().parent.parent / "reference"
    if repo_reference.parent.exists():
        return repo_reference

    return Path(os.path.expanduser("~/.config/opencode/templates"))


OUTPUT_DIR = resolve_output_dir()


def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_paragraph_spacing(style, before=0, after=0, line_spacing=None, line_rule=None):
    """Set paragraph spacing on a style."""
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing
    if line_rule:
        pf.line_spacing_rule = line_rule


def set_table_borders(table, color="DDDDDD", size="4"):
    """Set thin borders on all table edges."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_bottom_border_to_paragraph(paragraph, color="1A237E", size="6"):
    """Add a bottom border (rule) to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def configure_base_styles(doc, tight_spacing=False):
    """Configure common styles across all templates."""
    # --- Normal (body text) ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT_DARK
    line_sp = 1.15 if tight_spacing else 1.25
    set_paragraph_spacing(style, before=0, after=6, line_spacing=line_sp)

    # --- Heading 1 ---
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri Light'
    h1.font.size = Pt(16)
    h1.font.color.rgb = NAVY
    h1.font.bold = False
    set_paragraph_spacing(h1, before=18, after=6)

    # --- Heading 2 ---
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri Light'
    h2.font.size = Pt(13)
    h2.font.color.rgb = NAVY
    h2.font.bold = False
    set_paragraph_spacing(h2, before=14, after=4)

    # --- Heading 3 ---
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(11)
    h3.font.color.rgb = TEXT_DARK
    h3.font.bold = True
    set_paragraph_spacing(h3, before=10, after=3)

    # --- Heading 4 ---
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(11)
    h4.font.color.rgb = TEXT_GRAY
    h4.font.bold = True
    h4.font.italic = True
    set_paragraph_spacing(h4, before=8, after=2)

    # --- Title ---
    title = doc.styles['Title']
    title.font.name = 'Calibri Light'
    title.font.size = Pt(22)
    title.font.color.rgb = NAVY
    title.font.bold = False
    set_paragraph_spacing(title, before=0, after=4)

    # --- Subtitle ---
    subtitle = doc.styles['Subtitle']
    subtitle.font.name = 'Calibri Light'
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = BLUE
    subtitle.font.bold = False
    subtitle.font.italic = False
    set_paragraph_spacing(subtitle, before=0, after=12)

    # --- Quote / Block Quote ---
    quote = doc.styles['Quote']
    quote.font.name = 'Calibri'
    quote.font.size = Pt(11)
    quote.font.color.rgb = TEXT_GRAY
    quote.font.italic = True
    quote.paragraph_format.left_indent = Cm(1.0)
    set_paragraph_spacing(quote, before=6, after=6)

    # --- List Bullet ---
    lb = doc.styles['List Bullet']
    lb.font.name = 'Calibri'
    lb.font.size = Pt(11)
    lb.font.color.rgb = TEXT_DARK
    set_paragraph_spacing(lb, before=0, after=3)

    # --- List Number ---
    ln = doc.styles['List Number']
    ln.font.name = 'Calibri'
    ln.font.size = Pt(11)
    ln.font.color.rgb = TEXT_DARK
    set_paragraph_spacing(ln, before=0, after=3)

    return doc


def configure_page(doc, margin=1.0, header_text=None, footer_text=None):
    """Set page margins, optional header/footer."""
    section = doc.sections[0]
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

    if header_text:
        header = section.header
        p = header.paragraphs[0]
        p.text = header_text
        p.style.font.size = Pt(8)
        p.style.font.color.rgb = TEXT_GRAY
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = TEXT_GRAY

    if footer_text:
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = footer_text
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = TEXT_GRAY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_sample_table(doc, header_shading=True):
    """Add a sample table with professional styling."""
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set full width
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    # Borders
    set_table_borders(table, color="DDDDDD", size="4")

    # Header row
    headers = ["Column A", "Column B", "Column C"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
        if header_shading:
            set_cell_shading(cell, "1A237E")

    # Data rows with alternating shading
    for row_idx in range(1, 4):
        for col_idx in range(3):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = f"Data {row_idx}-{col_idx + 1}"
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = TEXT_DARK
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F5F8FC")

    return table


def add_code_block(doc):
    """Add a styled code block paragraph."""
    # Use 'No Spacing' as base and modify, or just add a raw paragraph
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run("def hello():\n    return 'world'")
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_DARK
    # Add shading to the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shd)
    # Add some padding via indentation
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    return p


# =============================================================================
# TEMPLATE 1: Customer Brief / Meeting Prep
# =============================================================================
def generate_brief_template():
    """Dense 1-2 page brief. No cover page. Starts with title then At a Glance table."""
    doc = Document()
    configure_base_styles(doc, tight_spacing=True)
    configure_page(doc, margin=0.85, footer_text="Confidential - Esper Inc.")

    # Title
    p = doc.add_paragraph("Customer Brief Title", style='Title')
    p.style.font.size = Pt(18)

    # Subtitle (date/context)
    p = doc.add_paragraph("Meeting context or date", style='Subtitle')

    # "At a Glance" summary table
    doc.add_paragraph("At a Glance", style='Heading 1')
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)
    set_table_borders(table, color="DDDDDD", size="4")

    labels = ["ARR", "Devices", "Stage", "Close Date"]
    values = ["$120K", "2,500", "Negotiation", "2026-07-15"]
    for i, (label, val) in enumerate(zip(labels, values)):
        cell = table.rows[0].cells[i]
        cell.text = label
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9)
        set_cell_shading(cell, "1A237E")
        table.rows[1].cells[i].text = val
        for run in table.rows[1].cells[i].paragraphs[0].runs:
            run.font.size = Pt(10)

    # Sections
    doc.add_paragraph("Background", style='Heading 1')
    doc.add_paragraph("Brief background on the customer relationship and current status.", style='Normal')

    doc.add_paragraph("Key Contacts", style='Heading 2')
    doc.add_paragraph("List of stakeholders and their roles.", style='Normal')

    doc.add_paragraph("Open Items", style='Heading 2')
    doc.add_paragraph("Current action items and blockers.", style='List Bullet')
    doc.add_paragraph("Second bullet point.", style='List Bullet')

    doc.add_paragraph("Talking Points", style='Heading 1')
    doc.add_paragraph("Prepared discussion topics for the meeting.", style='List Number')
    doc.add_paragraph("Second topic.", style='List Number')

    doc.add_paragraph("Recent Activity", style='Heading 1')
    add_sample_table(doc)

    path = OUTPUT_DIR / "esper-brief.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =============================================================================
# TEMPLATE 2: Technical Document / Analysis
# =============================================================================
def generate_technical_template():
    """Multi-page technical doc with TOC style, rule under H1, code blocks."""
    doc = Document()
    configure_base_styles(doc, tight_spacing=False)
    configure_page(doc, margin=1.0, footer_text="Esper Inc.")

    # Title
    doc.add_paragraph("Technical Document Title", style='Title')
    doc.add_paragraph("Esper Engineering", style='Subtitle')

    # TOC placeholder
    doc.add_paragraph("Table of Contents", style='Heading 1')
    doc.add_paragraph("[TOC will be generated by pandoc --toc]", style='Normal')

    doc.add_paragraph()  # spacer

    # Section with rule under H1
    p = doc.add_paragraph("Problem Statement", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")

    doc.add_paragraph(
        "Describe the technical problem being analyzed. Include relevant context, "
        "system constraints, and success criteria.",
        style='Normal'
    )

    doc.add_paragraph("Current Architecture", style='Heading 2')
    doc.add_paragraph(
        "Explanation of the current system design and its limitations.",
        style='Normal'
    )

    # Code block
    doc.add_paragraph("Implementation Details", style='Heading 3')
    add_code_block(doc)

    p = doc.add_paragraph("Analysis", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")

    doc.add_paragraph("Detailed technical analysis with supporting data.", style='Normal')

    # Table with minimal borders
    add_sample_table(doc)

    p = doc.add_paragraph("Recommendations", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")

    doc.add_paragraph("Recommended approach with justification.", style='List Number')
    doc.add_paragraph("Alternative approach considered.", style='List Number')

    doc.add_paragraph(
        "Important note or caveat about the analysis.",
        style='Quote'
    )

    path = OUTPUT_DIR / "esper-technical.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =============================================================================
# TEMPLATE 3: Executive Summary / One-Pager
# =============================================================================
def generate_executive_template():
    """Single page, high-impact. Large title, key metrics table, call-to-action."""
    doc = Document()
    configure_base_styles(doc, tight_spacing=True)
    configure_page(doc, margin=0.85)

    # Large title
    title_style = doc.styles['Title']
    title_style.font.size = Pt(24)
    doc.add_paragraph("Executive Summary", style='Title')
    doc.add_paragraph("Subtitle or customer name", style='Subtitle')

    # Key metrics — borderless table with bold labels
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)
    # No borders on this table
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    metrics = [("$450K ARR", "Revenue"), ("12,000", "Devices"), ("Q3 2026", "Target Close")]
    for i, (value, label) in enumerate(metrics):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value + "\n")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run2 = p.add_run(label)
        run2.font.size = Pt(9)
        run2.font.color.rgb = TEXT_GRAY

    doc.add_paragraph()  # spacer

    # Summary section
    doc.add_paragraph("Summary", style='Heading 1')
    doc.add_paragraph(
        "High-level overview of the situation, opportunity, or decision needed. "
        "This should be concise and executive-friendly.",
        style='Normal'
    )

    doc.add_paragraph("Key Points", style='Heading 2')
    doc.add_paragraph("First key insight or finding.", style='List Bullet')
    doc.add_paragraph("Second key insight or finding.", style='List Bullet')
    doc.add_paragraph("Third key insight or finding.", style='List Bullet')

    # Call to action box (light blue background paragraph)
    doc.add_paragraph("Recommendation", style='Heading 2')
    p = doc.add_paragraph()
    run = p.add_run("Next Steps: ")
    run.font.bold = True
    run.font.color.rgb = NAVY
    run2 = p.add_run("Clear call-to-action or decision request for the reader.")
    run2.font.color.rgb = TEXT_DARK
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F4FD" w:val="clear"/>')
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)

    path = OUTPUT_DIR / "esper-executive.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =============================================================================
# TEMPLATE 4: QBR / Report (Multi-page formal)
# =============================================================================
def generate_report_template():
    """Formal multi-page report with cover section, TOC, header, page numbers."""
    doc = Document()
    configure_base_styles(doc, tight_spacing=False)
    configure_page(doc, margin=1.0, header_text="Esper", footer_text="Confidential")

    # Cover section — big title centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run("Report Title")
    run.font.name = 'Calibri Light'
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared for Customer Name")
    run.font.name = 'Calibri Light'
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("June 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_GRAY

    # Page break after cover
    doc.add_page_break()

    # TOC placeholder
    doc.add_paragraph("Table of Contents", style='Heading 1')
    doc.add_paragraph("[Generated by pandoc --toc]", style='Normal')

    doc.add_page_break()

    # Content sections
    p = doc.add_paragraph("Executive Summary", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")
    doc.add_paragraph(
        "Brief overview of key findings and recommendations for this reporting period.",
        style='Normal'
    )

    p = doc.add_paragraph("Performance Metrics", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")

    doc.add_paragraph("Deployment Overview", style='Heading 2')
    doc.add_paragraph("Summary of current deployment state and changes.", style='Normal')

    # Sample metrics table
    add_sample_table(doc)

    doc.add_paragraph("Usage Trends", style='Heading 2')
    doc.add_paragraph("Analysis of feature adoption and usage patterns.", style='Normal')

    p = doc.add_paragraph("Recommendations", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")
    doc.add_paragraph("First recommendation with rationale.", style='List Number')
    doc.add_paragraph("Second recommendation with rationale.", style='List Number')

    p = doc.add_paragraph("Next Steps", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")
    doc.add_paragraph("Action item with owner and timeline.", style='List Bullet')
    doc.add_paragraph("Second action item.", style='List Bullet')

    path = OUTPUT_DIR / "esper-report.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =============================================================================
# TEMPLATE 5: Proposal / RFP Response
# =============================================================================
def generate_proposal_template():
    """Formal proposal with cover, numbered sections, requirement tables."""
    doc = Document()
    configure_base_styles(doc, tight_spacing=False)
    configure_page(doc, margin=1.0, footer_text="Confidential - Esper Inc.")

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    run = p.add_run("Proposal Title")
    run.font.name = 'Calibri Light'
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("In Response to Customer RFP")
    run.font.name = 'Calibri Light'
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    run = p.add_run("Submitted by: Esper\nDate: June 2026\nVersion: 1.0")
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_GRAY

    doc.add_page_break()

    # TOC
    doc.add_paragraph("Table of Contents", style='Heading 1')
    doc.add_paragraph("[Generated by pandoc --toc --number-sections]", style='Normal')

    doc.add_page_break()

    # Sections (these will be numbered by pandoc --number-sections)
    p = doc.add_paragraph("Company Overview", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")
    doc.add_paragraph(
        "Brief overview of Esper, our platform, and relevant experience.",
        style='Normal'
    )

    p = doc.add_paragraph("Technical Response", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")

    doc.add_paragraph("Platform Architecture", style='Heading 2')
    doc.add_paragraph("Description of the technical solution architecture.", style='Normal')

    doc.add_paragraph("Security & Compliance", style='Heading 2')
    doc.add_paragraph("Security posture, certifications, and compliance details.", style='Normal')

    # Requirements table (common in RFPs)
    doc.add_paragraph("Requirements Matrix", style='Heading 2')
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)
    set_table_borders(table, color="DDDDDD", size="4")

    headers = ["ID", "Requirement", "Response", "Status"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9)
        set_cell_shading(cell, "1A237E")

    sample_data = [
        ["REQ-001", "Must support MDM enrollment", "Fully supported via...", "Compliant"],
        ["REQ-002", "Must provide OTA updates", "Esper Airwave provides...", "Compliant"],
        ["REQ-003", "Must support geofencing", "Available via policy...", "Partial"],
    ]
    for row_idx, row_data in enumerate(sample_data, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F5F8FC")

    p = doc.add_paragraph("Implementation Plan", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")
    doc.add_paragraph("Phase 1: Discovery and setup (Weeks 1-2)", style='List Number')
    doc.add_paragraph("Phase 2: Pilot deployment (Weeks 3-4)", style='List Number')
    doc.add_paragraph("Phase 3: Full rollout (Weeks 5-8)", style='List Number')

    p = doc.add_paragraph("Pricing", style='Heading 1')
    add_bottom_border_to_paragraph(p, color=RULE_NAVY, size="4")
    doc.add_paragraph("Pricing details and commercial terms.", style='Normal')

    path = OUTPUT_DIR / "esper-proposal.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


# =============================================================================
# MAIN
# =============================================================================
if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print("Generating Esper pandoc reference templates...")
    print()
    generate_brief_template()
    generate_technical_template()
    generate_executive_template()
    generate_report_template()
    generate_proposal_template()
    print()
    print("Done. Templates saved to:", OUTPUT_DIR)
    print()
    print("Usage:")
    print("  pandoc input.md -o output.docx \\")
    print("    --reference-doc=~/.config/opencode/templates/esper-brief.docx \\")
    print("    --from=markdown+smart+pipe_tables+yaml_metadata_block \\")
    print("    --lua-filter=~/.opencode/skill/pandoc-docs/filters/strip-bookmarks.lua")
